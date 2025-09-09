import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def write_python_files_to_docx():
    # 定义文件路径和要处理的文件列表
    base_dir = r"E:\ZeroIG-main\ZeroIG-main"
    files_to_process = ["utils.py", "train.py", "test.py", 
                       "multi_read_data.py", "model.py", "loss.py"]
    output_file = r"E:\ZeroIG-main\ZeroIG-main\微光图像增强代码.docx"
    
    # 创建新的Word文档
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('微光图像增强代码', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 遍历所有文件
    for filename in files_to_process:
        file_path = os.path.join(base_dir, filename)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"警告：文件 {file_path} 不存在，跳过")
            continue
            
        # 添加文件名作为标题
        doc.add_heading(filename, level=1)
        
        # 读取文件内容
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
            # 添加代码段落
            code_para = doc.add_paragraph()
            code_para.add_run(content).font.name = 'Consolas'  # 使用等宽字体
            
            print(f"已处理: {filename}")
            
        except Exception as e:
            print(f"处理文件 {filename} 时出错: {str(e)}")
            # 添加错误信息到文档
            doc.add_paragraph(f"处理文件 {filename} 时出错: {str(e)}")
    
    # 保存文档
    doc.save(output_file)
    print(f"所有文件已成功写入 {output_file}")

if __name__ == "__main__":
    write_python_files_to_docx()
